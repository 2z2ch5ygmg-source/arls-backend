from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from email.message import EmailMessage
from html.parser import HTMLParser
import html
from io import BytesIO
import os
from pathlib import Path
import re
import shutil
import smtplib
import subprocess
import tempfile
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Mm
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph as RLParagraph
from reportlab.platypus import SimpleDocTemplate, Spacer
from reportlab.platypus import Table as RLTable
from reportlab.platypus import TableStyle
from reportlab.lib.styles import ParagraphStyle

from ..config import settings

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "employment_certificate.html"
_DEFAULT_TEMPLATE_CACHE: str | None = None

_KOREAN_FONT_NAME = "HYSMyeongJo-Medium"
_FALLBACK_FONT_NAME = "Helvetica"
_FONT_REGISTERED = False

EMPLOYMENT_CERTIFICATE_TEMPLATE_VARIABLES = (
    "company_name",
    "biz_reg_no",
    "ceo_name",
    "company_phone",
    "company_email",
    "company_address",
    "employee_name",
    "birth_date",
    "resident_no_masked",
    "org_name",
    "position_name",
    "hire_date",
    "issue_number",
    "issue_date",
    "issue_date_long",
    "purpose_label",
)

EMPLOYMENT_CERTIFICATE_TEMPLATE_REQUIRED_VARIABLES = (
    "company_name",
    "employee_name",
    "org_name",
    "hire_date",
    "issue_number",
    "issue_date",
    "purpose_label",
)

_TEMPLATE_PLACEHOLDER_ALIASES: dict[str, tuple[str, ...]] = {
    "company_name": ("company_name", "companyname", "회사명", "회사이름"),
    "biz_reg_no": ("biz_reg_no", "bizregno", "business_number", "사업자등록번호", "사업자등록번호"),
    "ceo_name": ("ceo_name", "대표자", "대표자명"),
    "company_phone": ("company_phone", "company_tel", "전화번호", "회사전화번호", "회사연락처"),
    "company_email": ("company_email", "email", "회사이메일", "이메일"),
    "company_address": ("company_address", "주소", "회사주소"),
    "employee_name": ("employee_name", "name", "성명", "이름", "직원명"),
    "birth_date": ("birth_date", "birthdate", "생년월일"),
    "org_name": ("org_name", "소속", "부서", "현장", "현장명"),
    "position_name": ("position_name", "position", "직급", "직위", "권한"),
    "hire_date": ("hire_date", "hiredate", "입사일", "채용일"),
    "issue_number": ("issue_number", "issue_no", "발급번호", "문서번호"),
    "issue_date": ("issue_date", "issued_at", "발급일"),
    "issue_date_long": ("issue_date_long", "발급일한글", "발급일자"),
    "purpose_label": ("purpose_label", "purpose", "용도", "제출용도", "발급용도"),
}


PURPOSE_LABEL_MAP = {
    "BANK": "은행 제출",
    "GOV": "관공서 제출",
    "CARD": "카드사 제출",
    "OTHER": "기타",
}

DOCX_LABEL_PATTERNS: dict[str, tuple[str, ...]] = {
    "employee_name": ("성명",),
    "resident_no_masked": ("주민번호", "주민등록번호"),
    "employee_address": ("주소",),
    "employee_phone": ("연락처", "전화번호", "연락처전화번호"),
    "company_name": ("회사명", "회사명소속"),
    "biz_reg_no": ("사업자", "사업자번호", "사업자등록번호"),
    "ceo_name": ("대표자", "대표자명"),
    "company_phone": ("대표전화", "대표연락처", "회사전화번호", "대표번호", "대표전화번호"),
    "company_address": ("소재지", "회사주소"),
    "employment_period": ("재직기간",),
    "purpose_label": ("발급용도",),
}

DOCX_PERSONAL_SECTION_TOKENS = ("인적사항",)
DOCX_EMPLOYMENT_SECTION_TOKENS = ("재직사항",)
DOCX_SECTION_HEADER_TOKENS = DOCX_PERSONAL_SECTION_TOKENS + DOCX_EMPLOYMENT_SECTION_TOKENS
DOCX_MARKER_FILLER_PATTERN = re.compile(r"^[-_~.=()\[\]\s□■▢▣▪▫◻◼◽◾]+$")


@dataclass
class MailResult:
    sent: bool
    error: str | None = None


@dataclass
class TemplateCell:
    text: str
    is_label: bool = False


@dataclass
class TemplateBlock:
    kind: str
    tag: str = "p"
    text: str = ""
    rows: list[list[TemplateCell]] | None = None


def _safe_text(value: Any, fallback: str = "-") -> str:
    text = str(value or "").strip()
    return text if text else fallback


def _resolve_paragraph_tag(style_name: str) -> str:
    normalized = str(style_name or "").strip().lower()
    if normalized.startswith("heading"):
        match = re.search(r"(\d+)", normalized)
        if match:
            level = max(1, min(6, int(match.group(1))))
            return f"h{level}"
        return "h2"
    return "p"


def _collect_docx_cell_text(cell) -> str:
    lines = []
    for paragraph in getattr(cell, "paragraphs", []) or []:
        text = str(paragraph.text or "").strip()
        if text:
            lines.append(text)
    merged = " ".join(lines).strip()
    return merged


def normalize_docx_label(text: str | None) -> str:
    value = str(text or "")
    value = value.replace("\xa0", " ")
    value = re.sub(r"\s+", "", value)
    value = re.sub(r"[:()\[\]{}<>\-–—·.,/\\\\]", "", value)
    return value.strip()


def _is_docx_value_placeholder(text: str | None) -> bool:
    raw = str(text or "").strip()
    if not raw:
        return True
    normalized = raw.replace("\xa0", " ").strip()
    if not normalized:
        return True
    if DOCX_MARKER_FILLER_PATTERN.fullmatch(normalized):
        return True
    return False


def _is_docx_section_header(text: str | None) -> bool:
    normalized = normalize_docx_label(text)
    if not normalized:
        return False
    return any(token in normalized for token in DOCX_SECTION_HEADER_TOKENS)


def _detect_docx_label_key(text: str | None) -> str | None:
    normalized = normalize_docx_label(text)
    if not normalized:
        return None
    for key, patterns in DOCX_LABEL_PATTERNS.items():
        for pattern in patterns:
            normalized_pattern = normalize_docx_label(pattern)
            if normalized == normalized_pattern or normalized.startswith(normalized_pattern):
                return key
    return None


def _resolve_docx_label_value(label_key: str, context: dict[str, Any], section: str | None) -> str:
    if label_key == "employee_address" and str(section or "").strip().lower() == "employment":
        return str(context.get("company_address") or "").strip()
    if label_key == "company_name":
        return str(context.get("company_name") or "").strip()
    if label_key == "employment_period":
        return str(context.get("employment_period") or "").strip()
    if label_key == "purpose_label":
        return str(context.get("purpose_label") or "").strip()
    return str(context.get(label_key) or "").strip()


def _set_docx_cell_text(cell, value: str) -> None:
    normalized_value = str(value or "").strip()
    lines = [line.strip() for line in normalized_value.splitlines()] if normalized_value else []
    if not lines:
        lines = [""]
    cell.text = lines[0]
    for line in lines[1:]:
        paragraph = cell.add_paragraph()
        paragraph.add_run(line)


def _iter_docx_paragraphs(document: DocumentObject) -> list[Paragraph]:
    paragraphs: list[Paragraph] = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def fill_employment_certificate_docx_by_labels(
    template_docx_bytes: bytes,
    context: dict[str, Any],
    *,
    seal_image_bytes: bytes | None = None,
) -> bytes:
    document: DocumentObject = Document(BytesIO(template_docx_bytes))

    for table in document.tables:
        section: str | None = "personal"
        for row in table.rows:
            row_cells = list(row.cells)
            if not row_cells:
                continue

            for scan_cell in row_cells:
                scan_normalized = normalize_docx_label(_collect_docx_cell_text(scan_cell))
                if not scan_normalized:
                    continue
                if any(token in scan_normalized for token in DOCX_PERSONAL_SECTION_TOKENS):
                    section = "personal"
                if any(token in scan_normalized for token in DOCX_EMPLOYMENT_SECTION_TOKENS):
                    section = "employment"

            for idx, label_cell in enumerate(row_cells):
                label_key = _detect_docx_label_key(_collect_docx_cell_text(label_cell))
                if not label_key:
                    continue
                value_text = _resolve_docx_label_value(label_key, context, section)
                if value_text is None:
                    continue

                candidate_index: int | None = None
                for value_idx in range(idx + 1, len(row_cells)):
                    value_cell = row_cells[value_idx]
                    value_cell_text = _collect_docx_cell_text(value_cell)
                    if _detect_docx_label_key(value_cell_text):
                        continue
                    if _is_docx_section_header(value_cell_text):
                        continue
                    if _is_docx_value_placeholder(value_cell_text):
                        candidate_index = value_idx
                        break
                    if candidate_index is None:
                        candidate_index = value_idx
                if candidate_index is None:
                    continue
                _set_docx_cell_text(row_cells[candidate_index], value_text)

    if seal_image_bytes:
        inserted = False
        paragraphs = _iter_docx_paragraphs(document)
        for paragraph in paragraphs:
            if not re.search(r"\(\s*인\s*\)", str(paragraph.text or "")):
                continue
            run = paragraph.add_run(" ")
            run.add_picture(BytesIO(seal_image_bytes), width=Mm(12))
            inserted = True
            break
        if not inserted:
            for paragraph in paragraphs:
                text = str(paragraph.text or "")
                if "대표이사" not in text and "대표자" not in text:
                    continue
                run = paragraph.add_run(" ")
                run.add_picture(BytesIO(seal_image_bytes), width=Mm(12))
                inserted = True
                break

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def fill_employment_certificate_docx_template(
    template_docx_bytes: bytes,
    context: dict[str, Any],
    *,
    seal_image_bytes: bytes | None = None,
) -> bytes:
    # Backward-compat alias: 기존 호출부는 이 함수를 사용해도 레이블 기반 채움 로직을 사용한다.
    return fill_employment_certificate_docx_by_labels(
        template_docx_bytes,
        context,
        seal_image_bytes=seal_image_bytes,
    )


def convert_docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    binary = shutil.which("libreoffice") or shutil.which("soffice")
    if not binary:
        raise RuntimeError("DOCX_PDF_CONVERTER_NOT_AVAILABLE")

    with tempfile.TemporaryDirectory(prefix="employment-certificate-") as temp_dir:
        temp_path = Path(temp_dir)
        source_path = temp_path / "employment_certificate.docx"
        source_path.write_bytes(docx_bytes)
        command = [
            binary,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(temp_path),
            str(source_path),
        ]
        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=90,
            env={
                **os.environ,
                **{"HOME": str(temp_path)},
            },
            check=False,
        )
        pdf_path = temp_path / "employment_certificate.pdf"
        if result.returncode != 0 or not pdf_path.exists():
            stderr = str(result.stderr or "").strip()
            stdout = str(result.stdout or "").strip()
            raise RuntimeError(f"DOCX_PDF_CONVERT_FAILED: {stderr or stdout or 'unknown'}")
        return pdf_path.read_bytes()


def convert_filled_docx_to_pdf(docx_bytes: bytes) -> bytes:
    # Backward-compat alias
    return convert_docx_to_pdf_bytes(docx_bytes)


def issue_employment_certificate_pdf_from_docx(
    template_docx_bytes: bytes,
    context: dict[str, Any],
    *,
    seal_image_bytes: bytes | None = None,
) -> bytes:
    filled_docx_bytes = fill_employment_certificate_docx_by_labels(
        template_docx_bytes,
        context,
        seal_image_bytes=seal_image_bytes,
    )
    return convert_docx_to_pdf_bytes(filled_docx_bytes)


def convert_docx_template_to_html(docx_bytes: bytes) -> str:
    document: DocumentObject = Document(BytesIO(docx_bytes))
    parts: list[str] = ['<div class="employment-certificate-template">']
    has_content = False

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            raw_text = str(paragraph.text or "").strip()
            if not raw_text:
                continue
            has_content = True
            tag = _resolve_paragraph_tag(getattr(getattr(paragraph, "style", None), "name", ""))
            safe_text = html.escape(raw_text).replace("\n", "<br/>")
            parts.append(f"<{tag}>{safe_text}</{tag}>")
            continue

        if isinstance(child, CT_Tbl):
            table = Table(child, document)
            has_content = True
            parts.append('<table class="certificate-table">')
            for row in table.rows:
                parts.append("<tr>")
                for cell in row.cells:
                    cell_text = _collect_docx_cell_text(cell)
                    safe_text = html.escape(cell_text).replace("\n", "<br/>") if cell_text else "&nbsp;"
                    parts.append(f"<td>{safe_text}</td>")
                parts.append("</tr>")
            parts.append("</table>")

    parts.append("</div>")
    if not has_content:
        raise ValueError("template docx is empty")
    return "\n".join(parts).strip()


def normalize_template_placeholders(template_html: str) -> str:
    normalized = str(template_html or "")
    if not normalized.strip():
        return normalized

    for canonical, aliases in _TEMPLATE_PLACEHOLDER_ALIASES.items():
        for alias in aliases:
            escaped = re.escape(alias)
            patterns = (
                re.compile(r"\{\{\s*" + escaped + r"\s*\}\}", flags=re.IGNORECASE),
                re.compile(r"\$\{\s*" + escaped + r"\s*\}", flags=re.IGNORECASE),
                re.compile(r"\[\[\s*" + escaped + r"\s*\]\]", flags=re.IGNORECASE),
                re.compile(r"<<\s*" + escaped + r"\s*>>", flags=re.IGNORECASE),
            )
            for pattern in patterns:
                normalized = pattern.sub(f"{{{{{canonical}}}}}", normalized)
    return normalized


def load_default_template_html() -> str:
    global _DEFAULT_TEMPLATE_CACHE
    if _DEFAULT_TEMPLATE_CACHE is not None:
        return _DEFAULT_TEMPLATE_CACHE
    _DEFAULT_TEMPLATE_CACHE = TEMPLATE_PATH.read_text(encoding="utf-8")
    return _DEFAULT_TEMPLATE_CACHE


def _register_pdf_font_once() -> str:
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return _KOREAN_FONT_NAME

    try:
        pdfmetrics.registerFont(UnicodeCIDFont(_KOREAN_FONT_NAME))
        _FONT_REGISTERED = True
        return _KOREAN_FONT_NAME
    except Exception:
        return _FALLBACK_FONT_NAME


def _wrap_text(text: str, max_chars: int = 48) -> list[str]:
    value = str(text or "").strip()
    if not value:
        return ["-"]

    words = re.split(r"(\s+)", value)
    lines: list[str] = []
    line = ""
    for token in words:
        if not token:
            continue
        if len(line + token) <= max_chars:
            line += token
            continue
        if line.strip():
            lines.append(line.strip())
            line = ""
        if len(token) > max_chars:
            chunk = token
            while len(chunk) > max_chars:
                lines.append(chunk[:max_chars])
                chunk = chunk[max_chars:]
            line = chunk
        else:
            line = token
    if line.strip():
        lines.append(line.strip())
    return lines or ["-"]


def build_issue_number(request_id: str, issued_at: datetime | None = None) -> str:
    now = issued_at or datetime.now(timezone.utc)
    token = re.sub(r"[^a-zA-Z0-9]", "", str(request_id or "")).upper()[:6] or "000000"
    return f"EC-{now:%Y%m%d}-{token}"


def render_employment_certificate_html(context: dict[str, Any], template_html: str | None = None) -> str:
    template = str(template_html or "").strip() or load_default_template_html()
    template = normalize_template_placeholders(template)
    # Never expose auto-injected debug helpers in final output PDF/HTML.
    template = re.sub(
        r"(?is)<section[^>]*class=\"[^\"]*hr-template-auto-fields[^\"]*\"[^>]*>.*?</section>",
        "",
        template,
    )
    template = re.sub(
        r"(?is)<div[^>]*class=\"[^\"]*hr-template-auto-fields[^\"]*\"[^>]*>.*?</div>",
        "",
        template,
    )

    rendered = template
    for key, raw in context.items():
        value = html.escape(str(raw or "").strip())
        rendered = re.sub(
            r"\{\{\s*" + re.escape(str(key)) + r"\s*\}\}",
            value,
            rendered,
            flags=re.IGNORECASE,
        )
    # Keep unresolved placeholders blank (not '-') to avoid polluted final certificate.
    rendered = re.sub(r"\{\{\s*[a-zA-Z0-9_]+\s*\}\}", "", rendered)
    return rendered


class _TemplateHtmlParser(HTMLParser):
    _PARA_TAGS = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "li"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=False)
        self.blocks: list[TemplateBlock] = []
        self._hidden_stack: list[bool] = []
        self._table_rows: list[list[TemplateCell]] | None = None
        self._current_row: list[TemplateCell] | None = None
        self._current_cell_text: list[str] | None = None
        self._current_cell_is_label = False
        self._current_para_tag: str | None = None
        self._current_para_text: list[str] = []

    def _is_hidden(self) -> bool:
        return bool(self._hidden_stack and self._hidden_stack[-1])

    def _append_text(self, value: str) -> None:
        if self._current_cell_text is not None:
            self._current_cell_text.append(value)
        elif self._current_para_tag is not None:
            self._current_para_text.append(value)

    def _flush_paragraph(self) -> None:
        if self._current_para_tag is None:
            return
        text = html.unescape("".join(self._current_para_text)).replace("\xa0", " ")
        text = re.sub(r"[ \t\r\f\v]+", " ", text)
        text = re.sub(r"\s*\n\s*", "\n", text).strip()
        if text:
            self.blocks.append(TemplateBlock(kind="paragraph", tag=self._current_para_tag, text=text))
        self._current_para_tag = None
        self._current_para_text = []

    def _start_paragraph(self, tag: str) -> None:
        self._flush_paragraph()
        self._current_para_tag = tag
        self._current_para_text = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = {str(k or "").lower(): str(v or "") for k, v in attrs}
        parent_hidden = self._is_hidden()
        style_text = attrs_dict.get("style", "").lower()
        class_text = attrs_dict.get("class", "").lower()
        is_hidden = parent_hidden or ("display:none" in style_text) or ("visibility:hidden" in style_text) or ("hidden" in attrs_dict)
        self._hidden_stack.append(is_hidden)

        if is_hidden:
            return

        tag = tag.lower()
        if tag in {"script", "style"}:
            self._hidden_stack[-1] = True
            return

        if tag == "br":
            self._append_text("\n")
            return

        if tag == "table":
            self._flush_paragraph()
            self._table_rows = []
            self._current_row = None
            return

        if self._table_rows is not None:
            if tag == "tr":
                self._current_row = []
                return
            if tag in {"td", "th"}:
                self._current_cell_text = []
                self._current_cell_is_label = "label" in class_text
                return
            return

        if tag in self._PARA_TAGS:
            self._start_paragraph(tag)
            return

        if tag == "div":
            if "title" in class_text:
                self._start_paragraph("h1")
            elif "section-title" in class_text:
                self._start_paragraph("h3")
            elif "issued-at" in class_text:
                self._start_paragraph("issued-at")
            elif "content" in class_text:
                self._start_paragraph("p")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        is_hidden = self._is_hidden()

        if not is_hidden:
            if self._table_rows is not None:
                if tag in {"td", "th"} and self._current_cell_text is not None:
                    text = html.unescape("".join(self._current_cell_text)).replace("\xa0", " ")
                    text = re.sub(r"[ \t\r\f\v]+", " ", text)
                    text = re.sub(r"\s*\n\s*", "\n", text).strip()
                    if self._current_row is not None:
                        self._current_row.append(TemplateCell(text=text, is_label=self._current_cell_is_label))
                    self._current_cell_text = None
                    self._current_cell_is_label = False
                elif tag == "tr":
                    if self._current_row:
                        self._table_rows.append(self._current_row)
                    self._current_row = None
                elif tag == "table":
                    if self._table_rows:
                        self.blocks.append(TemplateBlock(kind="table", tag="table", rows=self._table_rows))
                    self._table_rows = None
                    self._current_row = None
                    self._current_cell_text = None
                    self._current_cell_is_label = False
            else:
                if self._current_para_tag is not None:
                    if tag == self._current_para_tag:
                        self._flush_paragraph()
                    elif self._current_para_tag == "issued-at" and tag == "div":
                        self._flush_paragraph()

        if self._hidden_stack:
            self._hidden_stack.pop()

    def handle_data(self, data: str) -> None:
        if self._is_hidden():
            return
        if not data:
            return
        self._append_text(data)

    def close(self) -> None:
        super().close()
        self._flush_paragraph()


def _parse_template_blocks(rendered_html: str) -> list[TemplateBlock]:
    parser = _TemplateHtmlParser()
    parser.feed(str(rendered_html or ""))
    parser.close()
    return parser.blocks


def _build_paragraph_styles(font_name: str) -> dict[str, ParagraphStyle]:
    base = ParagraphStyle(
        "certificate-base",
        fontName=font_name,
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#111827"),
        spaceAfter=2,
    )
    return {
        "p": base,
        "li": ParagraphStyle("certificate-li", parent=base, leftIndent=12),
        "h1": ParagraphStyle("certificate-h1", parent=base, fontSize=24, leading=30, alignment=1, spaceAfter=10),
        "h2": ParagraphStyle("certificate-h2", parent=base, fontSize=18, leading=24, spaceBefore=4, spaceAfter=6),
        "h3": ParagraphStyle("certificate-h3", parent=base, fontSize=13, leading=18, spaceBefore=6, spaceAfter=4),
        "h4": ParagraphStyle("certificate-h4", parent=base, fontSize=12, leading=16, spaceBefore=4, spaceAfter=3),
        "h5": ParagraphStyle("certificate-h5", parent=base, fontSize=11, leading=15),
        "h6": ParagraphStyle("certificate-h6", parent=base, fontSize=10, leading=14),
        "issued-at": ParagraphStyle("certificate-issued-at", parent=base, alignment=2, spaceBefore=10, spaceAfter=4),
    }


def _to_paragraph_html(text: str) -> str:
    safe = html.escape(str(text or "").strip() or "-")
    safe = safe.replace("\n", "<br/>")
    return safe


def _contains_seal_marker(text: str) -> bool:
    normalized = str(text or "")
    if not normalized:
        return False
    if re.search(r"\(\s*인\s*\)", normalized):
        return True
    if re.search(r"\{\{\s*seal_image\s*\}\}", normalized, flags=re.IGNORECASE):
        return True
    if re.search(r"\[\s*seal\s*\]", normalized, flags=re.IGNORECASE):
        return True
    return False


def _make_seal_image_flowable(seal_image_bytes: bytes) -> RLImage | None:
    if not seal_image_bytes:
        return None
    try:
        stream = BytesIO(seal_image_bytes)
        image = RLImage(stream)
        image.drawWidth = 18 * mm
        image.drawHeight = 18 * mm
        image.hAlign = "RIGHT"
        return image
    except Exception:
        return None


def _build_table_flowable(
    rows: list[list[TemplateCell]],
    *,
    page_width: float,
    styles: dict[str, ParagraphStyle],
    seal_image_bytes: bytes | None,
) -> tuple[RLTable, bool]:
    col_count = max((len(row) for row in rows), default=1)
    if col_count <= 0:
        col_count = 1

    available = page_width
    if col_count == 4:
        col_widths = [available * 0.18, available * 0.32, available * 0.18, available * 0.32]
    elif col_count == 2:
        col_widths = [available * 0.26, available * 0.74]
    elif col_count == 3:
        col_widths = [available * 0.24, available * 0.38, available * 0.38]
    else:
        col_widths = [available / col_count for _ in range(col_count)]

    seal_inserted = False
    table_data: list[list[Any]] = []
    label_cells: list[tuple[int, int]] = []

    for r_idx, row in enumerate(rows):
        output_row: list[Any] = []
        normalized_row = row + [TemplateCell(text="", is_label=False)] * max(0, col_count - len(row))
        for c_idx, cell in enumerate(normalized_row[:col_count]):
            cell_text = str(cell.text or "").strip()
            if cell.is_label:
                label_cells.append((r_idx, c_idx))

            if seal_image_bytes and _contains_seal_marker(cell_text):
                seal_inserted = True
                clean_text = re.sub(r"\{\{\s*seal_image\s*\}\}", "", cell_text, flags=re.IGNORECASE)
                clean_text = re.sub(r"\[\s*seal\s*\]", "", clean_text, flags=re.IGNORECASE)
                content: list[Any] = []
                if clean_text.strip():
                    content.append(RLParagraph(_to_paragraph_html(clean_text), styles["p"]))
                    content.append(Spacer(1, 2))
                seal_image = _make_seal_image_flowable(seal_image_bytes)
                if seal_image is not None:
                    content.append(seal_image)
                output_row.append(content if content else RLParagraph("&nbsp;", styles["p"]))
                continue

            paragraph = RLParagraph(_to_paragraph_html(cell_text or "-"), styles["p"])
            output_row.append(paragraph)
        table_data.append(output_row)

    table = RLTable(table_data, colWidths=col_widths, repeatRows=0)
    style_commands: list[tuple] = [
        ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#D0D5DD")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]

    if label_cells:
        for r_idx, c_idx in label_cells:
            style_commands.append(("BACKGROUND", (c_idx, r_idx), (c_idx, r_idx), colors.HexColor("#F2F4F7")))
    elif col_count == 4:
        for r_idx in range(len(table_data)):
            style_commands.append(("BACKGROUND", (0, r_idx), (0, r_idx), colors.HexColor("#F2F4F7")))
            style_commands.append(("BACKGROUND", (2, r_idx), (2, r_idx), colors.HexColor("#F2F4F7")))

    table.setStyle(TableStyle(style_commands))
    return table, seal_inserted


def _generate_pdf_from_html_layout(
    rendered_html: str,
    *,
    seal_image_bytes: bytes | None = None,
) -> bytes:
    font_name = _register_pdf_font_once()
    paragraph_styles = _build_paragraph_styles(font_name)

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=14 * mm,
    )

    blocks = _parse_template_blocks(rendered_html)
    story: list[Any] = []
    page_inner_width = A4[0] - document.leftMargin - document.rightMargin
    seal_inserted = False

    for block in blocks:
        if block.kind == "table" and block.rows:
            table, inserted = _build_table_flowable(
                block.rows,
                page_width=page_inner_width,
                styles=paragraph_styles,
                seal_image_bytes=seal_image_bytes,
            )
            seal_inserted = seal_inserted or inserted
            story.append(table)
            story.append(Spacer(1, 6))
            continue

        if block.kind == "paragraph":
            text = str(block.text or "").strip()
            if not text:
                continue

            style_key = block.tag if block.tag in paragraph_styles else "p"
            if seal_image_bytes and _contains_seal_marker(text):
                clean_text = re.sub(r"\{\{\s*seal_image\s*\}\}", "", text, flags=re.IGNORECASE)
                clean_text = re.sub(r"\[\s*seal\s*\]", "", clean_text, flags=re.IGNORECASE)
                if clean_text.strip():
                    story.append(RLParagraph(_to_paragraph_html(clean_text), paragraph_styles.get(style_key, paragraph_styles["p"])))
                seal_image = _make_seal_image_flowable(seal_image_bytes)
                if seal_image is not None:
                    seal_inserted = True
                    story.append(seal_image)
                    story.append(Spacer(1, 4))
                continue

            story.append(RLParagraph(_to_paragraph_html(text), paragraph_styles.get(style_key, paragraph_styles["p"])))
            if style_key in {"h1", "h2", "h3", "h4", "h5", "h6", "issued-at"}:
                story.append(Spacer(1, 4))
            continue

    if seal_image_bytes and not seal_inserted:
        seal_image = _make_seal_image_flowable(seal_image_bytes)
        if seal_image is not None:
            signature = RLTable(
                [[RLParagraph("(인)", paragraph_styles["p"]), seal_image]],
                colWidths=[12 * mm, 24 * mm],
                hAlign="RIGHT",
            )
            signature.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("ALIGN", (0, 0), (-1, -1), "RIGHT"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ]
                )
            )
            story.append(Spacer(1, 8))
            story.append(signature)

    if not story:
        story.append(RLParagraph("재직증명서", paragraph_styles["h2"]))

    document.build(story)
    return buffer.getvalue()


def _generate_plain_text_fallback_pdf(
    context: dict[str, Any],
    *,
    seal_image_bytes: bytes | None = None,
    rendered_html: str,
) -> bytes:
    plain_text = rendered_html
    plain_text = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", plain_text)
    plain_text = re.sub(r"(?i)<br\s*/?>", "\n", plain_text)
    plain_text = re.sub(r"(?i)</(p|div|tr|h1|h2|h3|h4|h5|h6|li|table|section|article)>", "\n", plain_text)
    plain_text = re.sub(r"(?i)<li[^>]*>", "• ", plain_text)
    plain_text = re.sub(r"<[^>]+>", " ", plain_text)
    plain_text = html.unescape(plain_text)
    plain_lines = [line.strip() for line in plain_text.splitlines() if line.strip()]
    if not plain_lines:
        plain_lines = ["재직증명서"]

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    page_width, page_height = A4

    font_name = _register_pdf_font_once()

    def draw_label(x: float, y: float, text: str, size: int = 10):
        pdf.setFont(font_name, size)
        pdf.drawString(x, y, _safe_text(text, "-"))

    y = page_height - 56

    pdf.setFont(font_name, 24)
    pdf.drawCentredString(page_width / 2, y, "재직증명서")
    y -= 30

    pdf.setLineWidth(1)
    pdf.line(48, y, page_width - 48, y)
    y -= 20

    for line in plain_lines:
        wrapped = _wrap_text(line, 68)
        for item in wrapped:
            draw_label(52, y, item, 10)
            y -= 14
            if y < 120:
                pdf.showPage()
                pdf.setFont(font_name, 10)
                y = page_height - 56

    if seal_image_bytes:
        try:
            image = ImageReader(BytesIO(seal_image_bytes))
            pdf.drawImage(
                image,
                page_width - 160,
                max(48, y - 42),
                width=90,
                height=90,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            draw_label(page_width - 150, y - 20, "(도장 이미지 로드 실패)", 8)

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def generate_employment_certificate_pdf(
    context: dict[str, Any],
    *,
    seal_image_bytes: bytes | None = None,
    rendered_html: str | None = None,
    template_html: str | None = None,
) -> bytes:
    resolved_html = str(rendered_html or "").strip()
    if not resolved_html:
        resolved_html = render_employment_certificate_html(context, template_html=template_html)

    try:
        return _generate_pdf_from_html_layout(
            resolved_html,
            seal_image_bytes=seal_image_bytes,
        )
    except Exception:
        return _generate_plain_text_fallback_pdf(
            context,
            seal_image_bytes=seal_image_bytes,
            rendered_html=resolved_html,
        )


def _smtp_client():
    if settings.smtp_ssl:
        return smtplib.SMTP_SSL(
            settings.smtp_host,
            settings.smtp_port,
            timeout=settings.smtp_timeout_seconds,
        )
    return smtplib.SMTP(
        settings.smtp_host,
        settings.smtp_port,
        timeout=settings.smtp_timeout_seconds,
    )


def send_certificate_mail(
    *,
    to_email: str,
    subject: str,
    html_body: str,
    attachment_name: str,
    attachment_bytes: bytes,
) -> MailResult:
    target = str(to_email or "").strip()
    if not target:
        return MailResult(sent=False, error="RECIPIENT_EMPTY")

    if not settings.mail_enabled:
        return MailResult(sent=False, error="MAIL_DISABLED")
    if not settings.smtp_host:
        return MailResult(sent=False, error="SMTP_HOST_EMPTY")

    message = EmailMessage()
    message["Subject"] = f"{settings.mail_subject_prefix} {subject}".strip()
    message["From"] = settings.mail_from
    message["To"] = target
    message.set_content("재직증명서가 발급되었습니다. 첨부파일을 확인해 주세요.")
    message.add_alternative(html_body, subtype="html")
    message.add_attachment(
        attachment_bytes,
        maintype="application",
        subtype="pdf",
        filename=attachment_name,
    )

    try:
        with _smtp_client() as client:
            if not settings.smtp_ssl and settings.smtp_starttls:
                client.starttls()
            if settings.smtp_username:
                client.login(settings.smtp_username, settings.smtp_password)
            client.send_message(message)
    except Exception as exc:
        return MailResult(sent=False, error=str(exc))

    return MailResult(sent=True, error=None)


def build_purpose_label(purpose_code: str, purpose_text: str | None) -> str:
    normalized_code = str(purpose_code or "").strip().upper()
    label = PURPOSE_LABEL_MAP.get(normalized_code, normalized_code or "기타")
    if normalized_code == "OTHER" and str(purpose_text or "").strip():
        return f"기타 - {str(purpose_text).strip()}"
    return label
