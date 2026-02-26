from __future__ import annotations

from difflib import SequenceMatcher
import re
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from ..utils.address_norm import normalize_address_text

DATE_PATTERN = re.compile(r"(?P<year>\d{2,4})\s*[./-]\s*(?P<month>\d{1,2})\s*[./-]\s*(?P<day>\d{1,2})")
PHONE_PATTERN = re.compile(r"(01[016789])[-\s]?(\d{3,4})[-\s]?(\d{4})")
MGMT_NO_PATTERN = re.compile(r"[A-Za-z0-9]+(?:[-_][A-Za-z0-9]+)?")
SITE_HINT_SPLIT_PATTERN = re.compile(r"[\r\n/|,;]+")
NUMBER_TOKEN_RE = re.compile(r"^\d+[a-z가-힣]?$", flags=re.IGNORECASE)
ROAD_TOKEN_RE = re.compile(r".*(대로|로|길)$")
DISTRICT_TOKEN_RE = re.compile(r".*(시|군|구)$")
KOREAN_KEYWORD_RE = re.compile(r"[가-힣]{2,}")
SITE_NAME_NOISE_TOKENS = {"애플", "스토어", "현장", "매장", "센터", "점"}

FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    # 경비원명부 라벨 고정 파서 (요구사항 기준 9개)
    "management_no": ("관리번호", "관리 번호"),
    "name": ("성명", "성 명"),
    "address": ("주소",),
    "placement_text": ("배치지", "배치 지"),
    "birthdate": ("생년월일", "생년 월일"),
    "phone": ("전화번호", "전화 번호"),
    "training_cert_no": ("경비원 신임교육 이수증 교부번호", "교부번호", "교부 번호"),
    "hire_date": ("채용일", "채용 일"),
    "leave_date": ("퇴직일", "퇴직 일"),
}

IMAGE_EXT_TO_MIME = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".bmp": "image/bmp",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


def _clean_text(value: str | None) -> str:
    text = str(value or "").replace("\u3000", " ").strip()
    return re.sub(r"\s+", " ", text).strip()


def _normalize_label(value: str | None) -> str:
    text = _clean_text(value)
    # DOCX 고정 라벨 매칭용: 한글은 그대로 유지하고 공백/줄바꿈/구분 문장부호만 제거
    return re.sub(r"[\s\r\n\t:：()\[\]{}<>·]+", "", text)


def _field_alias_keys() -> set[str]:
    keys: set[str] = set()
    for aliases in FIELD_ALIASES.values():
        for alias in aliases:
            normalized = _normalize_label(alias)
            if normalized:
                keys.add(normalized)
    return keys


def _collect_docx_lines(document: Document) -> list[str]:
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        for line in re.split(r"[\r\n]+", text):
            normalized = _clean_text(line)
            if normalized:
                lines.append(normalized)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = _clean_text(cell.text)
                if not text:
                    continue
                for line in re.split(r"[\r\n]+", text):
                    normalized = _clean_text(line)
                    if normalized:
                        lines.append(normalized)

    return lines


def _collect_docx_pairs(document: Document) -> dict[str, str]:
    alias_keys = _field_alias_keys()
    pairs: dict[str, str] = {}

    # "라벨: 값" 형태가 문단/셀 한 줄에 있는 경우 우선 수집
    for line in _collect_docx_lines(document):
        cleaned_line = _clean_text(line)
        if not cleaned_line:
            continue
        for delimiter in (":", "："):
            if delimiter not in cleaned_line:
                continue
            left, right = cleaned_line.split(delimiter, 1)
            key = _normalize_label(left)
            val = _clean_text(right)
            if key in alias_keys and val and key not in pairs:
                pairs[key] = val

    for table in document.tables:
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            if not any(cells):
                continue

            for idx, raw_label in enumerate(cells):
                key = _normalize_label(raw_label)
                if key not in alias_keys:
                    continue
                if key in pairs and _clean_text(pairs.get(key)):
                    continue

                candidate = ""
                if idx + 1 < len(cells):
                    candidate = _clean_text(cells[idx + 1])
                # 다음 셀이 또 라벨이면 다다음 셀 fallback
                if candidate and _normalize_label(candidate) in alias_keys:
                    candidate = ""
                if not candidate and idx + 2 < len(cells):
                    fallback = _clean_text(cells[idx + 2])
                    if fallback and _normalize_label(fallback) not in alias_keys:
                        candidate = fallback
                if candidate:
                    pairs[key] = candidate

    return pairs


def _pick_field_value(pairs: dict[str, str], lines: list[str], aliases: tuple[str, ...]) -> str:
    alias_keys = [_normalize_label(alias) for alias in aliases]

    for alias in alias_keys:
        if not alias:
            continue
        for key, value in pairs.items():
            if alias in key or key in alias:
                candidate = _clean_text(value)
                if candidate:
                    return candidate

    for line in lines:
        for alias in aliases:
            pattern = rf"{re.escape(alias)}\s*[:：]?\s*(.+)$"
            match = re.search(pattern, line, flags=re.IGNORECASE)
            if match:
                candidate = _clean_text(match.group(1))
                if candidate:
                    return candidate
    return ""


def _normalize_date(value: str | None) -> str:
    text = _clean_text(value)
    if not text:
        return ""
    match = DATE_PATTERN.search(text)
    if not match:
        return ""
    year = int(match.group("year"))
    month = int(match.group("month"))
    day = int(match.group("day"))
    if year < 100:
        year = year + 2000 if year < 50 else year + 1900
    if not (1 <= month <= 12 and 1 <= day <= 31):
        return ""
    return f"{year:04d}-{month:02d}-{day:02d}"


def _normalize_phone(value: str | None) -> str:
    text = _clean_text(value)
    if not text:
        return ""
    match = PHONE_PATTERN.search(text)
    if match:
        return f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
    digits = re.sub(r"\D+", "", text)
    if len(digits) in (10, 11):
        if len(digits) == 10:
            return f"{digits[:3]}-{digits[3:6]}-{digits[6:]}"
        return f"{digits[:3]}-{digits[3:7]}-{digits[7:]}"
    return text


def _normalize_management_no(value: str | None) -> str:
    text = _clean_text(value)
    if not text:
        return ""
    match = MGMT_NO_PATTERN.search(text.replace(" ", ""))
    if match:
        return match.group(0)
    return text.replace(" ", "")


def parse_guard_roster_docx(docx_bytes: bytes) -> dict[str, str]:
    document = Document(BytesIO(docx_bytes))
    lines = _collect_docx_lines(document)
    pairs = _collect_docx_pairs(document)

    extracted: dict[str, str] = {}
    for field_name, aliases in FIELD_ALIASES.items():
        extracted[field_name] = _pick_field_value(pairs, lines, aliases)

    extracted["management_no"] = _normalize_management_no(extracted.get("management_no"))
    # management_no 문자열을 그대로 보존(leading zero 유지)하기 위한 명시 필드
    extracted["management_no_str"] = extracted["management_no"]
    extracted["name"] = _clean_text(extracted.get("name"))
    extracted["birthdate"] = _normalize_date(extracted.get("birthdate"))
    extracted["hire_date"] = _normalize_date(extracted.get("hire_date"))
    extracted["leave_date"] = _normalize_date(extracted.get("leave_date"))
    extracted["phone"] = _normalize_phone(extracted.get("phone"))
    extracted["address"] = _clean_text(extracted.get("address"))
    extracted["placement_text"] = _clean_text(extracted.get("placement_text"))
    extracted["training_cert_no"] = _clean_text(extracted.get("training_cert_no"))
    return extracted


def extract_primary_docx_photo(docx_bytes: bytes) -> tuple[bytes | None, str | None, str | None]:
    try:
        with ZipFile(BytesIO(docx_bytes)) as archive:
            images: list[tuple[int, bytes, str, str]] = []
            for info in archive.infolist():
                filename = str(info.filename or "")
                if not filename.startswith("word/media/"):
                    continue
                ext = Path(filename).suffix.lower()
                if ext not in IMAGE_EXT_TO_MIME:
                    continue
                data = archive.read(info)
                if not data:
                    continue
                images.append((len(data), data, IMAGE_EXT_TO_MIME[ext], Path(filename).name))
    except (BadZipFile, KeyError):
        return None, None, None

    if not images:
        return None, None, None
    images.sort(key=lambda item: item[0], reverse=True)
    _, payload, mime, filename = images[0]
    return payload, mime, filename


def _build_site_hint_text(*, placement_text: str, address_text: str) -> str:
    fragments: list[str] = []
    for raw in (_clean_text(placement_text), _clean_text(address_text)):
        if not raw:
            continue
        for chunk in SITE_HINT_SPLIT_PATTERN.split(raw):
            normalized = _clean_text(chunk)
            if normalized:
                fragments.append(normalized)
    if not fragments:
        return ""
    # 순서를 유지한 중복 제거
    deduped: list[str] = []
    seen: set[str] = set()
    for fragment in fragments:
        key = fragment.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(fragment)
    return " ".join(deduped)


def _sequence_similarity(a: str, b: str) -> float:
    left = str(a or "").strip()
    right = str(b or "").strip()
    if not left or not right:
        return 0.0
    return float(SequenceMatcher(None, left, right).ratio())


def _tokenize_address(value: str) -> list[str]:
    text = str(value or "").strip()
    if not text:
        return []
    tokens: list[str] = []
    for token in text.split():
        normalized = token.strip()
        if normalized:
            tokens.append(normalized)
    return tokens


def _road_number_city_score(query_norm: str, address_norm: str) -> tuple[float, bool, bool, bool]:
    query_tokens = set(_tokenize_address(query_norm))
    address_tokens = set(_tokenize_address(address_norm))
    if not query_tokens or not address_tokens:
        return 0.0, False, False, False

    query_road = {token for token in query_tokens if ROAD_TOKEN_RE.match(token)}
    addr_road = {token for token in address_tokens if ROAD_TOKEN_RE.match(token)}
    road_matched = bool(query_road & addr_road)

    query_num = {token for token in query_tokens if NUMBER_TOKEN_RE.match(token)}
    addr_num = {token for token in address_tokens if NUMBER_TOKEN_RE.match(token)}
    number_matched = bool(query_num & addr_num)

    query_district = {token for token in query_tokens if DISTRICT_TOKEN_RE.match(token)}
    addr_district = {token for token in address_tokens if DISTRICT_TOKEN_RE.match(token)}
    district_matched = bool(query_district & addr_district)

    boost = 0.0
    if road_matched:
        boost += 0.2
    if number_matched:
        boost += 0.35
    if district_matched:
        boost += 0.1
    return boost, road_matched, number_matched, district_matched


def _extract_site_name_keywords(site_name: str) -> set[str]:
    raw_tokens = KOREAN_KEYWORD_RE.findall(str(site_name or ""))
    keywords: set[str] = set()
    for token in raw_tokens:
        normalized = _clean_text(token)
        if len(normalized) < 2:
            continue
        if normalized in SITE_NAME_NOISE_TOKENS:
            continue
        keywords.add(normalized)
    return keywords


def _extract_query_keywords(query_text: str) -> set[str]:
    raw_tokens = KOREAN_KEYWORD_RE.findall(str(query_text or ""))
    keywords: set[str] = set()
    for token in raw_tokens:
        normalized = _clean_text(token)
        if len(normalized) < 2:
            continue
        keywords.add(normalized)
    return keywords


def _keyword_match_bonus(query_text: str, site_name: str) -> tuple[float, list[str]]:
    query_keywords = _extract_query_keywords(query_text)
    site_keywords = _extract_site_name_keywords(site_name)
    if not query_keywords or not site_keywords:
        return 0.0, []

    matched: list[str] = []
    for keyword in site_keywords:
        if keyword in query_keywords:
            matched.append(keyword)
            continue
        if keyword and keyword in query_text:
            matched.append(keyword)
            continue
        if any(keyword in token or token in keyword for token in query_keywords):
            matched.append(keyword)

    matched = sorted(set(matched))
    if len(matched) >= 2:
        return 0.45, matched
    if len(matched) >= 1:
        return 0.35, matched
    return 0.0, []


def _query_contamination_penalty(query_text: str, strong_keyword_match: bool) -> tuple[float, list[str]]:
    if strong_keyword_match:
        return 0.0, []

    penalty = 0.0
    reasons: list[str] = []
    clean = _clean_text(query_text)
    district_tokens = set(re.findall(r"[가-힣]{2,}구", clean))
    if len(district_tokens) >= 2:
        penalty += 0.15
        reasons.append("MULTI_DISTRICT")

    lowered = clean.lower()
    if "서울" in clean and "경기도" in clean:
        penalty += 0.15
        reasons.append("MIXED_REGION")
    elif "seoul" in lowered and "gyeonggi" in lowered:
        penalty += 0.15
        reasons.append("MIXED_REGION")

    return penalty, reasons


def _site_match_score(*, query_norm: str, site_name: str, address_norm: str) -> float:
    if not query_norm:
        return 0.0
    _ = site_name  # signature compatibility (site_name is used by caller in keyword stage)
    string_similarity = _sequence_similarity(query_norm, address_norm)
    address_component, _, _, _ = _road_number_city_score(query_norm, address_norm)
    # 문자열 유사도는 최대 0.15까지만 반영
    return round(min(1.0, address_component + (string_similarity * 0.15)), 4)


def select_site_match_query_text(*, placement_text: str, address_text: str) -> str:
    placement_clean = _clean_text(placement_text)
    if len(placement_clean) >= 6:
        return placement_clean
    return _clean_text(address_text)


def match_site_candidates(
    *,
    placement_text: str,
    address_text: str = "",
    sites: list[dict],
    threshold: float = 0.60,
    top_n: int = 3,
) -> dict[str, object]:
    # 근무현장(배치지)을 우선 사용하고, 배치지가 충분하지 않을 때만 주소를 fallback으로 사용한다.
    query_text = select_site_match_query_text(
        placement_text=placement_text,
        address_text=address_text,
    )
    query_norm = normalize_address_text(query_text)

    ranked: list[dict[str, object]] = []
    for site in sites:
        site_code = _clean_text(site.get("site_id") or site.get("site_code"))
        site_name = _clean_text(site.get("site_name"))
        site_address = _clean_text(site.get("address_text") or site.get("address"))
        site_address_norm = _clean_text(site.get("address_norm")) or normalize_address_text(site_address)
        string_similarity = _sequence_similarity(query_norm, site_address_norm)
        keyword_bonus, matched_keywords = _keyword_match_bonus(query_text, site_name)
        address_score = _site_match_score(
            query_norm=query_norm,
            site_name=site_name,
            address_norm=site_address_norm,
        )
        address_boost, road_matched, number_matched, district_matched = _road_number_city_score(query_norm, site_address_norm)
        strong_keyword = keyword_bonus >= 0.45
        contamination_penalty, penalty_reasons = _query_contamination_penalty(query_text, strong_keyword)
        score = max(0.0, min(1.0, keyword_bonus + address_score - contamination_penalty))
        ranked.append(
            {
                "site_id": site_code,
                "site_code": site_code,
                "site_name": site_name or site_code,
                "score": round(float(score), 3),
                "keyword_bonus": round(float(keyword_bonus), 3),
                "matched_keywords": matched_keywords,
                "address_score": round(float(address_score), 3),
                "string_similarity": round(float(string_similarity), 3),
                "road_boost": 0.2 if road_matched else 0.0,
                "number_boost": 0.35 if number_matched else 0.0,
                "district_boost": 0.1 if district_matched else 0.0,
                "address_boost_total": round(float(address_boost), 3),
                "contamination_penalty": round(float(contamination_penalty), 3),
                "penalty_reasons": penalty_reasons,
            }
        )

    ranked.sort(key=lambda item: float(item.get("score") or 0), reverse=True)
    candidates = ranked[: max(1, int(top_n or 3))]
    best = candidates[0] if candidates else None
    best_score = float(best.get("score") or 0) if best else 0.0
    second_score = float(candidates[1].get("score") or 0) if len(candidates) > 1 else 0.0
    score_gap = round(max(0.0, best_score - second_score), 3)
    auto_site_code = str(best.get("site_code") or "").strip() if best else ""
    auto_site_name = str(best.get("site_name") or "").strip() if best else ""
    best_candidate = f"{auto_site_code}/{auto_site_name}".strip("/") if (auto_site_code or auto_site_name) else ""
    threshold_score = max(0.65, float(threshold))
    gap_auto = best_score >= 0.55 and score_gap >= 0.15
    threshold_auto = best_score >= threshold_score
    should_auto = bool(auto_site_code) and (threshold_auto or gap_auto)

    print(
        "AUTO_MATCH:",
        query_norm,
        best_candidate,
        f"best={round(best_score, 3)}",
        f"second={round(second_score, 3)}",
        f"gap={score_gap}",
        f"threshold_auto={threshold_auto}",
        f"gap_auto={gap_auto}",
    )

    status = "READY" if should_auto else "NEEDS_SITE_PICK"
    if not query_norm:
        match_reason = "QUERY_EMPTY"
    elif not sites:
        match_reason = "INDEX_EMPTY"
    elif not candidates:
        match_reason = "NO_CANDIDATE"
    elif status == "READY" and gap_auto and not threshold_auto:
        match_reason = "AUTO_MATCHED_BY_GAP"
    elif status == "READY":
        match_reason = "AUTO_MATCHED"
    elif best_score < 0.55 or score_gap < 0.15:
        match_reason = "LOW_CONFIDENCE"
    else:
        match_reason = "NEEDS_REVIEW"
    return {
        "site_code": auto_site_code if status == "READY" else "",
        "site_name": auto_site_name if status == "READY" else "",
        "confidence": round(best_score, 3),
        "second_score": round(second_score, 3),
        "score_gap": score_gap,
        "candidates": candidates,
        "status": status,
        "query_norm": query_norm,
        "match_reason": match_reason,
    }


def build_employee_code_from_management_no(site_code: str, management_no: str) -> str:
    left = _clean_text(site_code)
    right = _clean_text(management_no)
    if not left or not right:
        return ""
    return f"{left}-{right}"
